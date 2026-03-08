"""Generate FEATURE_STATUS_REPORT.docx from FEATURE_STATUS_REPORT.md."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_row(row, hex_color="D9E1F2"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, size=9, bold=False, color=None, align=None):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    h.paragraph_format.space_after = Pt(4)
    return h


def add_para(doc, text, size=10, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_code_block(doc, text):
    for line in text.strip().split("\n"):
        p = doc.add_paragraph(style="No Spacing")
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8)
    doc.add_paragraph()


def make_table(doc, headers, rows, col_widths=None, severity_col=None):
    """Create a styled table.

    severity_col: if set, column index to apply severity-based coloring
                  (CRITICAL=red bg, WARNING=yellow bg, INFO=blue bg)
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    shade_row(hdr_row, "2F5496")
    for i, h in enumerate(headers):
        set_cell_text(hdr_row.cells[i], h, size=9, bold=True, color=(0xFF, 0xFF, 0xFF))

    # Data rows
    SEVERITY_COLORS = {
        "CRITICAL": "FDE0D9",
        "CRIT": "FDE0D9",
        "WARNING": "FFF2CC",
        "WARN": "FFF2CC",
        "WARN / CRIT": "FFF2CC",
        "WARNING / CRITICAL": "FFF2CC",
        "INFO": "D9E8FC",
    }
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        # Alternate row shading
        if ri % 2 == 0:
            shade_row(row, "F2F5FA")
        for ci, cell_text in enumerate(row_data):
            set_cell_text(row.cells[ci], cell_text, size=9)
        # Severity coloring
        if severity_col is not None and severity_col < len(row_data):
            sev_text = str(row_data[severity_col]).strip()
            if sev_text in SEVERITY_COLORS:
                shade_row(row, SEVERITY_COLORS[sev_text])

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def make_kv_table(doc, pairs, key_width=2.0, val_width=4.0, key_shade="2F5496"):
    """Create a key-value style table (2 columns, key has dark bg)."""
    table = doc.add_table(rows=len(pairs), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    for i, (k, v) in enumerate(pairs):
        shade_cell(table.rows[i].cells[0], "E8ECF4")
        set_cell_text(table.rows[i].cells[0], k, size=9, bold=True)
        set_cell_text(table.rows[i].cells[1], v, size=9)
        table.rows[i].cells[0].width = Inches(key_width)
        table.rows[i].cells[1].width = Inches(val_width)

    doc.add_paragraph()
    return table


def make_dashboard(doc, items, cols=4):
    """Create a dashboard-style metrics table."""
    n_rows = (len(items) + cols - 1) // cols
    table = doc.add_table(rows=n_rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(items):
        r, c = divmod(i, cols)
        cell = table.rows[r].cells[c]
        shade_cell(cell, "2F5496")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(value)
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(4)
        run2 = p2.add_run(label)
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0xBB, 0xCC, 0xEE)

    # Fill remaining cells
    for i in range(len(items), n_rows * cols):
        r, c = divmod(i, cols)
        shade_cell(table.rows[r].cells[c], "2F5496")

    doc.add_paragraph()
    return table


def build_docx():
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Default font
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10)

    # ══════════════════════════════════════════════════════════════════════
    # Title
    # ══════════════════════════════════════════════════════════════════════
    title = doc.add_heading("KooDynaErrorAnalyzer", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("종합 기능 보고서  v0.2.1")
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("최종 갱신: 2026-03-08")
    mr.font.size = Pt(9)
    mr.font.italic = True
    doc.add_paragraph()

    # Dashboard
    make_dashboard(doc, [
        ("소스 코드", "10,100줄"),
        ("소스 파일", "36개"),
        ("에러 코드 DB", "68개"),
        ("검증 케이스", "2,440개"),
        ("파서 모듈", "11개"),
        ("분석 엔진", "10개"),
        ("데이터 모델", "27개"),
        ("SVG 차트", "8종"),
    ], cols=4)

    # ══════════════════════════════════════════════════════════════════════
    # 1. 프로젝트 개요
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. 프로젝트 개요", 1)
    add_para(doc, (
        "KooDynaErrorAnalyzer는 LS-DYNA MPP(Massively Parallel Processing) 유한요소 "
        "시뮬레이션 결과를 자동 파싱·분석하여 수치 불안정, 성능 병목, 실패 원인을 "
        "한국어로 진단하는 CLI/GUI 도구입니다. "
        "실무 엔지니어가 수십~수백 MB 규모의 d3hsp, glstat, matsum 등 텍스트 출력을 "
        "직접 열어 보는 대신, 몇 초 만에 핵심 문제를 CRITICAL / WARNING / INFO 등급으로 "
        "분류하고 LS-DYNA 키워드 수준의 구체적 해결 방안을 제시합니다."
    ))

    make_kv_table(doc, [
        ("언어", "Python 3.10+"),
        ("총 코드 규모", "10,100줄 / 36개 소스 파일"),
        ("데이터 모델", "27개 dataclass + 2개 Enum (Severity, TerminationStatus)"),
        ("에러/경고 코드 DB", "68개 코드 (한국어 FEM 이론 설명 + 해결 방안 포함)"),
        ("출력 형식", "한국어 터미널(Rich) / HTML(독립형) / JSON(구조화)"),
        ("분석 대상 파일", "d3hsp, glstat, mesXXXX, matsum, rcforc, nodout, bndout, status.out, "
                       "load_profile.csv, cont_profile.csv, slurm .err, 입력 덱(.k/.dyn)"),
        ("빌드", "PyInstaller 단일 실행파일 (Linux x86-64 / Windows x64)"),
        ("검증 규모", "/data 내 2,495개 디렉터리 인덱싱 → 2,440개 분석 완료 (97.8%)"),
    ])

    add_heading(doc, "실행 명령 예시", 2)
    make_table(doc,
        ["명령", "설명"],
        [
            ["koodyna <결과폴더>/", "단일 폴더 분석 (터미널 한국어 리포트)"],
            ["koodyna <결과폴더>/ --html report.html", "HTML 리포트 생성 (브라우저 자동 오픈)"],
            ["koodyna <결과폴더>/ -o report.json", "JSON 리포트 생성 (후처리·자동화용)"],
            ["koodyna --scan /data", "/data 전체 재귀 탐색 → 인덱스 갱신"],
            ["koodyna --analyze --output-dir /reports", "인덱스 pending 항목 배치 분석"],
            ["koodyna --list-index", "인덱스 현황 테이블 출력"],
        ],
        col_widths=[3.0, 3.5]
    )

    # ══════════════════════════════════════════════════════════════════════
    # 2. 아키텍처
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. 아키텍처 — 6-Phase 분석 파이프라인", 1)

    add_heading(doc, "2.1 소스 구조", 2)
    make_table(doc,
        ["디렉터리 / 파일", "모듈 수", "역할"],
        [
            ["__main__.py", "1", "진입점"],
            ["cli.py", "1", "CLI (argparse) — 단일/배치/GUI/스캔 모드"],
            ["analyzer.py", "1", "오케스트레이터 (6-Phase 파이프라인)"],
            ["models.py", "1", "데이터 모델 (24개 dataclass, 2개 Enum)"],
            ["scanner.py", "1", "디렉터리 인덱싱 + 배치 분석"],
            ["parsers/", "11", "d3hsp, glstat, messag, matsum, rcforc, nodout, bndout, slurm, element_mapper, profile, status"],
            ["analysis/", "10", "diagnostics, energy, timestep, warnings, contact, performance, numerical_instability, failure_analysis, matsum_analysis, implicit_diagnostics"],
            ["knowledge/", "1", "error_db (68개 에러/경고 코드 DB)"],
            ["report/", "4", "terminal(Rich), html_report(독립형+SVG 차트), json_report(구조화), svg_chart(순수 SVG 생성)"],
        ],
        col_widths=[1.8, 0.8, 3.9]
    )

    add_heading(doc, "2.2 파이프라인 단계", 2)
    make_table(doc,
        ["Phase", "단계명", "역할", "주요 모듈"],
        [
            ["1", "파일 탐색", "d3hsp, glstat, mes*, matsum, nodout, bndout, slurm.err, 입력 덱 자동 탐지", "analyzer._discover_files()"],
            ["2", "파싱", "10개 파서로 구조화된 데이터 추출", "parsers/*"],
            ["3", "리포트 조립", "파싱 결과 → Report 모델 통합 + element_mapper 파트 보완", "analyzer.run()"],
            ["4", "분석", "에너지·타임스텝·경고·접촉·성능·matsum 분석", "analysis/*"],
            ["5", "진단", "수치 불안정·실패 원인·Implicit solver·Slurm 장애 진단", "diagnostics.py 등"],
            ["6", "후처리", "중복/과민 Finding 제거 (deduplication)", "_deduplicate_findings()"],
        ],
        col_widths=[0.5, 1.0, 2.8, 2.2]
    )

    # ══════════════════════════════════════════════════════════════════════
    # 3. 파서
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. 파서 모듈 (11개)", 1)

    add_heading(doc, "3.1 핵심 파서 (3개)", 2)
    make_table(doc,
        ["파서", "대상 파일", "추출 정보"],
        [
            ["d3hsp", "d3hsp", "시뮬레이션 헤더, 모델 크기, 종료 상태, 경고/에러 카운트, 키워드 카운트, "
             "타임스텝, 파트 정의, 성능 타이밍, 접촉 정의, MPP 프로세서, 에너지, 질량 속성, 도메인 분해 메트릭"],
            ["glstat", "glstat", "사이클별 에너지 스냅샷 (KE, IE, HG, Sliding, External work, 에너지 비율), "
             "NaN/Inf 감지, dt2ms 포맷 지원, 제어 요소 추적"],
            ["messag", "mes0000~mesNNNN", "MPP 전 랭크 경고/에러 카운트, 에러 상세(요소 번호), "
             "초기 관통, 인터페이스 경고 요약, 최소 타임스텝, 서프스 타임스텝, 접촉 dt 상한"],
        ],
        col_widths=[0.8, 1.5, 4.2]
    )

    add_heading(doc, "3.2 보조 파서 (8개)", 2)
    make_table(doc,
        ["파서", "대상 파일", "추출 정보", "신규"],
        [
            ["matsum", "matsum", "재료별 에너지 시계열: IE, KE, HG, x/y/z momentum, eroded energy", ""],
            ["rcforc", "rcforc", "접촉 인터페이스별 반력·모멘트 시계열 (SURFA/SURFB), 레전드 파싱, single surface 감지", "v0.2.1"],
            ["nodout", "nodout", "노드별 변위·속도 시계열. Shooting node / 고주파 진동 감지", ""],
            ["bndout", "bndout", "경계 반력·모멘트 시계열. 반력 스파이크·진동 감지", ""],
            ["slurm", "slurm-*.err", "Slurm HPC 잡 에러. Segfault·MPI 에러·exit code·스택 트레이스", ""],
            ["element_mapper", "*.k / *.dyn", "입력 덱에서 *ELEMENT_SOLID/SHELL/BEAM → 요소 ID→파트 ID 매핑", "연동"],
            ["profile", "load_profile.csv, cont_profile.csv", "프로세서별 컴포넌트 부하 프로파일", ""],
            ["status", "status.out", "CPU/clock per zone-cycle 타이밍, 잔여 시간 예측", ""],
        ],
        col_widths=[1.0, 1.5, 3.0, 0.7]
    )

    add_heading(doc, "3.3 파서 기술 특징", 2)
    make_table(doc,
        ["특징", "설명"],
        [
            ["스트리밍 상태 머신", "d3hsp: 7-상태 FSM (HEADER→KEYWORD_COUNTS→CONTROL_INFO→PART_DEFS→CONTACTS→BODY→TAIL). 146K+ 줄 파일을 한 줄씩 처리하여 메모리 효율적"],
            ["NaN/Inf 안전 파싱", "glstat에서 NaN, Inf, ***** 등 비정상 값을 안전하게 감지·기록. 비정상 에너지 상태 즉시 플래그"],
            ["MPP 다중 파일 병합", "mes0000~mesNNNN 모든 랭크의 경고/에러를 병합. 비-제로 랭크에서만 발생하는 에러도 포착"],
            ["matsum 재료별 추적", "각 material ID별 독립 시계열 구성. 재료명 자동 추출. Hourglass 에너지 추적"],
            ["rcforc 접촉 반력 파싱", "SURFA/SURFB(또는 slave/master) 양면 반력·모멘트 추출. 레전드 섹션에서 인터페이스 타이틀 매핑"],
            ["입력 덱 요소 매핑", "element_mapper: *ELEMENT_SOLID/SHELL/BEAM 파싱으로 요소→파트 매핑. 타임스텝 제어 요소의 파트 식별률 향상"],
        ],
        col_widths=[1.8, 4.7]
    )

    # ══════════════════════════════════════════════════════════════════════
    # 4. 분석 엔진
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. 분석 엔진 (10개 모듈)", 1)

    add_heading(doc, "4.1 수치 불안정 분석 (numerical_instability.py)", 2)
    add_para(doc, "nodout / bndout / glstat / slurm 데이터를 기반으로 수치적 건전성을 종합 진단합니다. 커버리지 약 95%.", size=9)
    make_table(doc,
        ["#", "진단 항목", "임계값", "심각도", "소스"],
        [
            ["1", "Shooting node", "|v| > 1,000 m/s", "CRITICAL", "nodout"],
            ["2", "고주파 진동", "ZCR > 10 kHz", "WARNING", "nodout"],
            ["3", "반력 스파이크", "max/mean > 100", "CRITICAL", "bndout"],
            ["4", "반력 진동", "ZCR 기반", "WARNING", "bndout"],
            ["5", "Hourglass 에너지 지배", "HG/IE > 10% / 50%", "WARN / CRIT", "glstat"],
            ["6", "운동 에너지 폭발", "100x 급증", "CRITICAL", "glstat"],
            ["7", "KE/IE 비율 이상", "> 10 (준정적 해석)", "WARNING", "glstat"],
            ["8", "접촉 슬라이딩 과다", "Slide/IE > 30%", "WARNING", "glstat"],
            ["9", "접촉 에너지 스파이크", "50x 급증", "CRITICAL", "glstat"],
            ["10", "타임스텝 변동", "10x 급락", "WARNING", "glstat"],
            ["11", "NaN 에너지 감지", "NaN/Inf in glstat", "CRITICAL", "glstat"],
            ["12", "음수 에너지 성분", "IE < 0 또는 Slide < 0", "CRITICAL", "glstat"],
            ["13", "Slurm 장애 진단", "segfault, MPI error, exit code", "CRITICAL", "slurm"],
        ],
        col_widths=[0.3, 1.6, 1.7, 1.0, 0.7],
        severity_col=3,
    )

    add_heading(doc, "4.2 에너지 분석 (energy.py)", 2)
    make_table(doc,
        ["진단 항목", "임계값", "심각도"],
        [
            ["에너지 보존 편차", "ratio > 1.05 또는 < 0.95", "WARNING"],
            ["에너지 보존 심각 위반", "ratio > 1.10 또는 < 0.90", "CRITICAL"],
            ["에너지 발산", "총 에너지 5% 이상 성장", "CRITICAL"],
            ["Hourglass 비율", "> 10% / > 50%", "WARNING / CRITICAL"],
            ["Sliding 에너지 과다", "> 15% (total 대비)", "WARNING"],
        ],
        col_widths=[2.2, 2.2, 1.6],
        severity_col=2,
    )
    add_para(doc, "음수 sliding CRITICAL 존재 시 Sliding WARNING 자동 억제 (deduplication).", size=9, italic=True)

    add_heading(doc, "4.3 타임스텝 분석 (timestep.py)", 2)
    make_table(doc,
        ["기능", "설명"],
        [
            ["최소 타임스텝 추출", "100개 최소 dt 수집 + 제어 파트 식별"],
            ["타임스텝 붕괴 감지", "dt < 1E-11 + Warning 40509 → CRITICAL (시뮬레이션 실질 정지)"],
            ["파트별 지배율 계산", "특정 파트가 최소 dt의 50% 이상 차지 → WARNING"],
            ["element_mapper 연동", "요소 ID → 파트 ID 보완 (d3hsp에 없는 경우 입력 덱에서 매핑)"],
        ],
        col_widths=[2.0, 4.5]
    )

    add_heading(doc, "4.4 성능 분석 (performance.py)", 2)
    make_table(doc,
        ["진단 항목", "임계값", "심각도"],
        [
            ["접촉 계산 과다", "> 40% / > 50%", "WARNING / CRITICAL"],
            ["Force gather 과다", "> 5% / > 10%", "WARNING / CRITICAL"],
            ["Mass scaling 과다", "> 5%", "WARNING"],
            ["MPP 부하 불균형", "> 15% / >= 100%", "WARNING / CRITICAL"],
            ["MPI 스케일링 예측", "Amdahl 법칙 기반", "INFO"],
        ],
        col_widths=[2.2, 2.2, 1.6],
        severity_col=2,
    )
    add_para(doc, "부하 균형 양호 시 INFO를 발생시키지 않음 — 노이즈 방지 원칙 적용.", size=9, italic=True)

    add_heading(doc, "4.5 재료별 Hourglass 분석 (matsum_analysis.py) — v0.2.0 신규", 2)
    add_para(doc, (
        "matsum 파일에서 각 재료(material)별 Hourglass 에너지 / 내부 에너지 비율을 계산합니다. "
        "글로벌 HG/IE 비율이 낮아도 특정 재료/파트에 hourglass가 집중될 수 있습니다."
    ))
    make_table(doc,
        ["진단 항목", "임계값", "심각도", "의미"],
        [
            ["재료별 HG/IE", "> 20%", "CRITICAL", "해당 재료 파트의 ELFORM 변경 또는 HG control 강화 필요"],
            ["재료별 HG/IE", "10~20%", "WARNING", "모니터링 필요, 메시 품질 점검 권장"],
        ],
        col_widths=[1.5, 1.0, 1.0, 3.0],
        severity_col=2,
    )
    add_para(doc, "HTML 리포트에 재료별 HG 테이블 섹션 추가. CRITICAL/WARNING 재료는 행 배경색으로 강조.", size=9, italic=True)

    add_heading(doc, "4.6 Implicit Solver 진단 (implicit_diagnostics.py) — v0.2.0 신규", 2)
    add_para(doc, (
        "*CONTROL_IMPLICIT_* 키워드 존재 여부로 implicit 해석을 자동 감지합니다. "
        "암묵적 시간 적분 특유의 수렴·안정성 문제를 진단하며, "
        "각 에러에 대해 LS-DYNA 키워드 수준의 해결 방안을 제시합니다."
    ))
    make_table(doc,
        ["진단 항목", "에러 코드", "심각도", "원인 설명"],
        [
            ["강성 행렬 특이", "Error 60004", "CRITICAL", "K 행렬 역행렬 불가 — 구속 불충분, 좌굴, 재료 softening"],
            ["Line search 실패", "Error 60303", "CRITICAL", "에너지 감소 스텝 크기 미발견 — 급격한 비선형성"],
            ["Newton-Raphson 발산", "Error 60315", "CRITICAL", "잔차 증가 — 구조적 불안정, 접촉 상태 급변"],
            ["수렴 속도 저하", "Warning 60121", "WARNING", "과다 반복 — 하중 스텝 또는 수렴 기준 조정 필요"],
            ["에너지 수지 불균형", "(비율 > 2.0)", "CRITICAL", "Newmark-beta/HHT-alpha 안정 조건 위반"],
            ["정상 수행", "—", "INFO", "Implicit 감지, 수렴 오류 없이 완료"],
        ],
        col_widths=[1.5, 1.2, 1.0, 2.8],
        severity_col=2,
    )

    add_heading(doc, "4.7 경고·에러 분석 (warnings.py)", 2)
    make_table(doc,
        ["기능", "설명"],
        [
            ["코드별 조회", "error_db에서 코드별 심각도, 한국어 설명, FEM 이론, 권장사항 조회"],
            ["Tied contact 상세", "Warning 50135/50136 — 인터페이스별 상세 진단"],
            ["MPP 에러 병합", "mes 파일 에러를 d3hsp와 병합 (비-제로 랭크 에러 포함)"],
            ["Fallback", "미등록 코드: 코드 범위 기반 자동 심각도 판정"],
        ],
        col_widths=[1.8, 4.7]
    )

    add_heading(doc, "4.8 종합 진단 (diagnostics.py)", 2)
    make_table(doc,
        ["진단 항목", "조건", "심각도", "설명"],
        [
            ["INCOMPLETE: 초기화 미시작", "actual=0, target=0", "CRITICAL", "시뮬레이션 미시작 — 입력 파일 오류, 라이선스 실패"],
            ["INCOMPLETE: 첫 사이클 미도달", "actual=0, target>0", "CRITICAL", "초기화 후 즉시 실패 — 초기 관통, 재료 초기화 오류"],
            ["INCOMPLETE: 실행 중 크래시", "actual>0", "CRITICAL", "X% 진행 후 비정상 종료 — segfault, 수치 발산"],
            ["접촉 dt 안정성", "min_dt > contact_dt_limit", "WARNING", "접촉 dt가 요소 dt보다 큼 — 접촉 안정성 위험"],
            ["에너지 비율 폭주", "ratio > 4.0 / > 3.0", "CRIT / WARN", "에너지 수지 심각 불균형"],
            ["파트 타임스텝 지배", "특정 파트 > 50%", "WARNING", "해당 파트 메시 개선 또는 mass scaling 검토"],
        ],
        col_widths=[1.8, 1.5, 1.0, 2.2],
        severity_col=2,
    )
    add_para(doc, "INCOMPLETE 종료 상태 3분류: v0.2.0에서 기존 단일 메시지를 3가지 상황별로 세분화.", size=9, italic=True)

    add_heading(doc, "4.9 기타 분석 모듈", 2)
    make_table(doc,
        ["모듈", "기능"],
        [
            ["contact.py", "인터페이스별 CPU/Clock 비용 분석, 병목 인터페이스 식별"],
            ["failure_analysis.py", "messag + d3hsp 교차 분석, 실패 요소 → 파트 역추적"],
        ],
        col_widths=[2.0, 4.5]
    )

    add_heading(doc, "4.10 Finding 중복 제거 (deduplication)", 2)
    make_table(doc,
        ["억제 대상", "조건"],
        [
            ["High sliding interface energy (WARNING)", "Excessive contact sliding 또는 Negative sliding이 CRITICAL/WARNING으로 이미 존재"],
        ],
        col_widths=[3.0, 3.5]
    )

    # ══════════════════════════════════════════════════════════════════════
    # 5. 지식 베이스
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. 지식 베이스 — error_db (68개 코드)", 1)
    add_para(doc, (
        "LS-DYNA 에러·경고 코드 데이터베이스. 각 코드에 대해 한국어 FEM 이론적 배경 설명 + "
        "LS-DYNA 키워드 수준 해결 방안(4개 필드: severity, title, description, recommendation)을 포함합니다."
    ))

    add_heading(doc, "5.1 코드 범위별 등록 현황", 2)
    make_table(doc,
        ["범위", "카테고리", "등록 수", "주요 코드"],
        [
            ["10xxx~11xxx", "요소 에러", "7", "10100, 10103, 10133, 10246, 10305, 11507, 40100"],
            ["20xxx", "재료", "9", "20018, 20200, 20216, 20248, 20268, 20282, 20385, 20430, 20471, 20546, 20661"],
            ["21xxx", "곡선/테이블", "4", "21129, 21287, 21302, 21329"],
            ["30xxx", "초기화/접촉", "11", "30001, 30010, 30060, 30062, 30099, 30128, 30131, 30200, 30210, 30358, 30364, 30455"],
            ["40xxx", "런타임 경고", "18", "40003, 40004, 40024, 40455, 40456, 40509, 40515, 40532, 40533, 40534, 40538, 40540, 40552, 40565, 40571, 40864, 41200, 41213, 41234, 41314"],
            ["50xxx", "타이드 접촉", "4", "50120, 50134, 50135, 50136"],
            ["60xxx", "암시적 해석", "5", "60004, 60100, 60121, 60303, 60315"],
            ["70xxx~80xxx", "SPH/입자", "3", "70021, 70100, 80100"],
            ["90xxx", "라이선스", "1", "90001"],
            ["Fallback", "미등록 자동", "—", "코드 범위 기반 심각도 자동 판정"],
        ],
        col_widths=[1.1, 1.1, 0.6, 3.7]
    )

    add_heading(doc, "5.2 고빈도 코드 Top 18 (실측, /data 2,440개 케이스)", 2)
    make_table(doc,
        ["순위", "코드", "제목", "누적 발생"],
        [
            ["1", "40533", "Contact velocity too high (slave node)", "250,000+"],
            ["2", "40538", "Slave node released from contact (energy conserved)", "198,000+"],
            ["3", "21129", "Curve extrapolation beyond defined range", "102,000+"],
            ["4", "40532", "Slave node penetration velocity limit exceeded", "57,000+"],
            ["5", "50134", "Tied contact slave node not found", "8,651"],
            ["6", "30099", "Contact pair definition error", "3,654"],
            ["7", "40540", "Contact force limit exceeded", "1,158"],
            ["8", "21302", "Curve ID reference invalid or undefined", "320"],
            ["9", "40565", "Contact shooting node detected", "180"],
            ["10", "30060", "Part section initialization inconsistency", "144"],
            ["11", "40864", "Contact surface penetration limit", "91"],
            ["12", "30210", "Constrained node initialization warning", "78"],
            ["13", "40534", "Contact slave penetration depth exceeded", "68"],
            ["14", "21287", "Curve interpolation warning", "56"],
            ["15", "60303/60315", "Implicit solver 수렴 실패", "55"],
            ["16", "20661", "Material model convergence issue", "31"],
            ["17", "41234", "Contact interface removal threshold", "23"],
            ["18", "11507", "SPG particle stretch parameter error", "5"],
        ],
        col_widths=[0.4, 0.9, 3.5, 0.9]
    )

    add_heading(doc, "5.3 Fallback 심각도 판정", 2)
    make_table(doc,
        ["코드 범위", "자동 분류", "기본 심각도"],
        [
            ["10000~19999", "Element error", "CRITICAL"],
            ["20000~29999", "Material/Curve", "WARNING~CRITICAL"],
            ["30000~39999", "Initialization/Constraint", "WARNING"],
            ["40000~49999", "Runtime warning", "WARNING"],
            ["50000~59999", "Contact warning", "WARNING"],
            ["60000~69999", "Implicit solver", "CRITICAL"],
            ["70000+", "SPH/Particle", "WARNING"],
        ],
        col_widths=[1.5, 2.5, 2.5],
        severity_col=2,
    )

    # ══════════════════════════════════════════════════════════════════════
    # 6. 인덱싱 시스템
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. 결과 디렉터리 인덱싱 시스템 (scanner.py)", 1)
    add_para(doc, (
        "대규모 결과 디렉터리 트리를 재귀 탐색하여 d3plot이 있는 모든 폴더를 "
        "JSON 인덱스(~/.koodyna/index.json)로 관리합니다. "
        "분석 상태를 추적하여 중단 후 재시작도 안전하게 지원합니다."
    ))

    add_heading(doc, "6.1 scanner.py 주요 함수", 2)
    make_table(doc,
        ["함수", "역할"],
        [
            ["scan_for_result_dirs(base_dir)", "d3plot 재귀 탐색 → 결과 디렉터리 목록 반환"],
            ["update_index(index, dirs)", "신규 폴더만 추가 (기존 유지)"],
            ["mark_analyzed(index, dir, status)", "분석 상태 기록 (analyzed / failed)"],
            ["run_batch_scan(base_dir)", "탐색 + 인덱스 갱신"],
            ["run_batch_analyze(index_path, ...)", "pending 항목 배치 분석 실행 (50건 단위 중간 저장)"],
            ["print_index(index_path)", "인덱스 현황 테이블 출력"],
        ],
        col_widths=[2.5, 4.0]
    )

    add_heading(doc, "6.2 /data 스캔 실측 결과", 2)
    make_kv_table(doc, [
        ("탐색 기반 디렉터리", "/data"),
        ("총 인덱스 항목", "2,495개"),
        ("분석 완료 (analyzed)", "2,440개 (97.8%)"),
        ("미분석 (pending)", "54개"),
        ("건너뜀 (skipped)", "1개"),
        ("Study 타입 수", "38개 이상"),
        ("주요 Study 타입", "floor_wave_study(143), vapor_chamber(93), shield_can_forming(92), "
                          "pcb_vibration(88), rubber_advanced_study(56), battery_study, warpage_study, ball_drop_v3~v7 등"),
    ])

    add_heading(doc, "6.3 배치 분석 CLI 옵션", 2)
    make_table(doc,
        ["옵션", "설명", "예시"],
        [
            ["--scan <DIR>", "BASE_DIR 재귀 탐색 → 인덱스 갱신", "--scan /data"],
            ["--analyze", "인덱스 pending 항목 배치 분석", "--analyze --output-dir /reports"],
            ["--output-dir <DIR>", "리포트 저장 위치", "--output-dir /data/koodyna_reports"],
            ["--limit N", "최대 처리 건수 (0=전체)", "--limit 100"],
            ["--no-html", "JSON만 생성 (HTML 생략)", "--analyze --no-html"],
            ["--reanalyze", "완료된 항목도 재실행", "--analyze --reanalyze"],
            ["--list-index", "인덱스 현황 테이블 출력", "--list-index"],
            ["--index <PATH>", "인덱스 파일 경로 지정", "--index /home/user/my_idx.json"],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    add_heading(doc, "6.4 인덱스 상태값", 2)
    make_table(doc,
        ["상태", "의미"],
        [
            ["pending", "아직 분석되지 않은 디렉터리"],
            ["analyzed", "분석 완료 (JSON + HTML 생성됨)"],
            ["failed", "분석 시도 중 에러 발생"],
            ["skipped", "d3hsp/mes 파일 부재 등 분석 불가"],
        ],
        col_widths=[1.5, 5.0]
    )

    # ══════════════════════════════════════════════════════════════════════
    # 7. 리포트 시스템
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "7. 리포트 시스템 (4개 모듈, 18개 섹션)", 1)

    add_heading(doc, "7.1 출력 형식", 2)
    make_table(doc,
        ["형식", "모듈", "특징"],
        [
            ["터미널", "terminal.py", "Rich 라이브러리 기반 컬러 출력, 한국어 18개 섹션"],
            ["HTML", "html_report.py", "독립형 (CSS + SVG 차트 임베디드), 외부 의존성 없음, 브라우저 자동 오픈"],
            ["JSON", "json_report.py", "구조화 데이터, 배치 처리·비교 분석·후처리 자동화용"],
            ["SVG 차트", "svg_chart.py", "순수 SVG 시계열 그래프 — matplotlib 등 외부 의존성 없음"],
        ],
        col_widths=[0.8, 1.5, 4.2]
    )

    add_heading(doc, "7.2 리포트 18개 섹션", 2)
    make_table(doc,
        ["#", "섹션명", "내용", "비고"],
        [
            ["1", "시뮬레이션 헤더", "LS-DYNA 버전, 날짜, 호스트, 정밀도, 라이선스", ""],
            ["2", "모델 요약", "노드, 요소(solid/shell/beam/tshell/SPH), 파트, 접촉, SPC 수", ""],
            ["3", "종료 상태", "정상/에러/미완료, 목표 시간, 사이클 수, CPU/경과 시간", ""],
            ["4", "진단 결과", "CRITICAL / WARNING / INFO 분류 + 설명 + 권장사항", "핵심 섹션"],
            ["5", "경고/에러 요약", "코드별 횟수, 심각도, 관련 인터페이스", ""],
            ["6", "에너지 분석", "에너지 비율 범위, HG/IE 비율, Sliding/Total 비율", ""],
            ["7", "에너지 시계열 그래프", "SVG 4종: 에너지 성분, 외부일/접촉, 에너지 비율, 타임스텝", "v0.2.1"],
            ["8", "재료별 HG 에너지", "matsum 기반 재료별 HG/IE 비율 테이블", ""],
            ["9", "접촉 반력 그래프", "rcforc SVG: 상위 인터페이스 합력 + 개별 Fx/Fy/Fz", "v0.2.1"],
            ["10", "타임스텝 분석", "100개 최소 dt, 제어 파트, 타임스텝 히스토리", ""],
            ["11", "성능 타이밍", "컴포넌트별 CPU/Clock 비율 (바 차트)", ""],
            ["12", "접촉 타이밍", "인터페이스별 CPU/Clock 비용", ""],
            ["13", "MPP 부하 균형", "프로세서별 CPU 비율, 불균형도", ""],
            ["14", "MPI 스케일링", "Amdahl 법칙 기반 32/64/128/256코어 예측", ""],
            ["15", "접촉 서프스 dt", "인터페이스별 서페이스 타임스텝", ""],
            ["16", "파트 정의", "재료 타입, 밀도, 탄성계수, 포아송비, ELFORM", ""],
            ["17", "Slurm 잡 정보", "exit code, signal, MPI 에러, 스택 트레이스", ""],
            ["18", "질량 속성", "파트별 질량, 질량 중심, 관성 모멘트", ""],
        ],
        col_widths=[0.3, 1.5, 3.2, 0.8]
    )

    # ══════════════════════════════════════════════════════════════════════
    # 8. 검증
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "8. 검증 — 대규모 실증 분석", 1)

    add_heading(doc, "8.1 검증 규모", 2)
    make_kv_table(doc, [
        ("총 인덱스 디렉터리", "2,495개"),
        ("분석 완료", "2,440개 (97.8%)"),
        ("ERROR 종료 케이스", "약 200개"),
        ("INCOMPLETE 종료 케이스", "64개"),
        ("NORMAL 종료 케이스", "약 2,176개"),
        ("Study 타입", "38+ 종류"),
    ])

    add_heading(doc, "8.2 진단 커버리지", 2)
    make_table(doc,
        ["메트릭", "값", "의미"],
        [
            ["ERROR/INCOMPLETE → CRITICAL", "100%", "모든 실패 케이스에서 최소 1개 CRITICAL finding 발생"],
            ["WARNING+ 보유 비율", "95.8%", "전체 2,440개 케이스 중 95.8%에서 WARNING 이상 소견"],
            ["INFO-only 케이스", "93개 (3.8%)", "대부분 정상 완료된 implicit thermal 시뮬레이션"],
            ["Finding 없음", "10개 (0.4%)", "d3hsp만 있고 glstat/mes 없는 최소 구성 케이스"],
        ],
        col_widths=[2.3, 1.2, 3.0]
    )

    add_heading(doc, "8.3 감지된 실패 모드", 2)
    make_table(doc,
        ["실패 모드", "감지 빈도", "비고"],
        [
            ["에너지 발산 (에너지 비율 > 1.10)", "다수", "가장 빈번한 실패 패턴 중 하나"],
            ["Negative volume (Error 40509)", "다수", "가장 빈번한 에러 코드"],
            ["Segmentation fault (Signal 11)", "다수", "Slurm 파서로 감지"],
            ["NaN 검출 (Error 40455/40456)", "다수", "glstat NaN 감지 + 에러 코드"],
            ["타임스텝 붕괴 (dt < 1E-11)", "다수", "시뮬레이션 실질적 정지"],
            ["Implicit solver 수렴 실패", "55건", "Error 60303/60315"],
            ["재료별 HG 과다 (>20%)", "—", "matsum 보유 케이스에서 감지"],
            ["MPP 부하 불균형 >= 100%", "소수", "도메인 분해 문제"],
            ["MPI 통신 에러", "소수", "Slurm 에러 파일에서 감지"],
        ],
        col_widths=[2.5, 1.0, 3.0]
    )

    add_heading(doc, "8.4 초기 테스트 케이스 (10개 상세)", 2)
    make_table(doc,
        ["케이스", "종료", "CRIT", "WARN", "주요 실패 원인"],
        [
            ["results_normal", "정상", "0", "2", "기준 케이스 (HG 경고)"],
            ["ball_drop_v2_dp_ex09", "에러", "4", "3", "Negative volume (40509), 에너지 발산"],
            ["ball_drop_v2_dp_ex16", "에러", "6", "1", "NaN (40455/40456), 음수 Sliding"],
            ["ball_drop_v2_sp_ex06", "에러", "4", "3", "Negative volume, Slurm exit 3"],
            ["ball_drop_v2_sp_ex09", "미완료", "5", "0", "Segfault (Signal 11), 음수 Sliding"],
            ["ball_drop_v2_sp_ex14", "미완료", "5", "0", "Segfault, NaN 에너지"],
            ["ball_drop_v2_sp_ex16", "미완료", "2", "1", "Segfault, 에너지 편차"],
            ["ball_drop_v4_ex12", "에러", "4", "0", "Negative volume, MPP 불균형 129%"],
            ["level_study_ex06", "에러", "9", "1", "NaN + 타임스텝 붕괴 + 에너지 폭주"],
            ["level_study_ex08", "미완료", "4", "1", "KE 폭발, MPI 통신 에러"],
        ],
        col_widths=[2.0, 0.7, 0.5, 0.5, 2.8]
    )

    # ══════════════════════════════════════════════════════════════════════
    # 9. 노이즈 관리
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "9. 노이즈 관리", 1)
    add_para(doc, (
        "대규모 배치 분석(2,440건)에서 발견된 불필요한 INFO finding을 체계적으로 제거했습니다. "
        "분석 도구는 문제를 감지하는 것이 목적이며, 문제가 없는 항목을 보고하는 것은 노이즈입니다."
    ))
    make_table(doc,
        ["제거된 항목", "발생 빈도", "비율", "제거 이유"],
        [
            ["nodout/bndout 부재 안내 (INFO)", "4,455건", "93%", "대부분의 케이스에 해당하며 실질적 가치 없음"],
            ["MPP 부하 균형 양호 (INFO)", "1,332건", "55%", "문제 없는 상태를 보고할 필요 없음"],
        ],
        col_widths=[2.3, 1.0, 0.6, 2.6]
    )

    # ══════════════════════════════════════════════════════════════════════
    # 10. 빌드
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "10. 빌드 및 배포", 1)

    add_heading(doc, "10.1 빌드 스크립트", 2)
    make_table(doc,
        ["파일", "플랫폼", "용도"],
        [
            ["build_linux.sh", "Linux x86-64", "PyInstaller 단일 실행파일 빌드"],
            ["build_windows.bat", "Windows (CMD)", "PyInstaller 빌드"],
            ["build_windows.ps1", "Windows (PowerShell)", "PyInstaller 빌드"],
            ["install.sh", "Linux", "venv 생성 + 의존성(rich) 설치"],
            ["install.bat", "Windows", "venv 생성 + 의존성(rich) 설치"],
            ["koodyna.sh", "Linux", "실행 래퍼 스크립트"],
            ["koodyna.bat", "Windows", "실행 래퍼 스크립트"],
        ],
        col_widths=[2.0, 1.5, 3.0]
    )

    add_heading(doc, "10.2 의존성", 2)
    make_table(doc,
        ["패키지", "용도", "필수 여부"],
        [
            ["rich", "터미널 컬러 출력", "런타임 필수"],
            ["PyInstaller", "단일 실행파일 빌드", "빌드 시에만"],
            ["python-docx", "DOCX 보고서 생성", "문서 생성 시에만"],
        ],
        col_widths=[1.5, 3.0, 2.0]
    )
    add_para(doc, "Python 표준 라이브러리 외 최소 의존성으로 설계.", size=9, italic=True)

    # ══════════════════════════════════════════════════════════════════════
    # 11. 개발 이력
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "11. 개발 이력", 1)

    add_heading(doc, "v0.2.1 (2026-03-08, 현재 버전)", 2)
    make_table(doc,
        ["카테고리", "변경 내용"],
        [
            ["신규 모듈", "svg_chart.py — 순수 SVG 시계열 그래프 생성 (외부 의존성 없음, Tokyonight dark 테마)"],
            ["신규 모듈", "rcforc.py — 접촉 반력·모멘트 시계열 파서 (SURFA/SURFB, 레전드, single surface)"],
            ["HTML 차트", "에너지 시계열 SVG 4종: 에너지 성분(KE/IE/HG/Total), 외부일/접촉, 에너지 비율, dt 이력"],
            ["HTML 차트", "접촉 반력 SVG: 상위 10개 인터페이스 합력 요약 + 상위 5개 Fx/Fy/Fz 상세"],
            ["데이터 모델", "RcforcSnapshot, RcforcInterface, Series 추가 (models.py, rcforc.py, svg_chart.py)"],
        ],
        col_widths=[1.2, 5.3]
    )

    add_heading(doc, "v0.2.0 (2026-03-08)", 2)
    make_table(doc,
        ["카테고리", "변경 내용"],
        [
            ["신규 모듈", "matsum_analysis.py — 재료별 Hourglass 에너지/내부 에너지 비율 계산"],
            ["신규 모듈", "implicit_diagnostics.py — CONTROL_IMPLICIT_* 키워드 기반 implicit 해석 자동 감지, Error 60004/60303/60315 전문 진단"],
            ["파서 연동", "element_mapper: 입력 덱(.k/.dyn) 파싱으로 요소→파트 매핑 보완. 타임스텝 제어 요소의 파트 식별률 향상"],
            ["진단 개선", "INCOMPLETE 종료 상태 3분류: 초기화 실패 / 첫 사이클 크래시 / 실행 중 크래시 (진행률 % 표시)"],
            ["DB 확장", "error_db: 51개 → 68개 코드 (17개 신규, /data 실측 빈도 기반 우선순위 등록)"],
            ["노이즈 제거", "nodout/bndout INFO(4,455건), MPP 양호 INFO(1,332건) 제거"],
            ["HTML 개선", "matsum 재료별 HG 테이블 섹션 추가 + CRITICAL/WARNING 행 배경색 강조"],
            ["대규모 검증", "/data 2,495개 디렉터리 인덱싱, 2,440개 분석 완료, ERROR→CRITICAL 100% 커버리지 확인"],
        ],
        col_widths=[1.2, 5.3]
    )

    add_heading(doc, "v0.1.0 (2026-03-04)", 2)
    make_table(doc,
        ["카테고리", "변경 내용"],
        [
            ["신규 모듈", "scanner.py — /data 인덱싱 + 배치 분석 시스템"],
            ["DB 확장", "error_db 28개 코드 신규 등록 (실측 고빈도 기반)"],
            ["CLI 추가", "--scan / --analyze / --list-index / --reanalyze / --limit / --no-html"],
            ["진단 개선", "Finding 중복 제거 강화, MPP 부하 불균형 >= 100% → CRITICAL 승격"],
            ["진단 개선", "접촉 dt 안정성 진단: min_dt > contact_dt_limit일 때만 WARNING"],
        ],
        col_widths=[1.2, 5.3]
    )

    add_heading(doc, "v0.0.x (2026-03-03 이전)", 2)
    make_table(doc,
        ["카테고리", "변경 내용"],
        [
            ["핵심 파서", "d3hsp / glstat / messag 파서 완성"],
            ["보조 파서", "nodout / bndout / slurm 파서 추가"],
            ["분석 엔진", "수치 불안정 진단 95% 커버리지 달성"],
            ["분석 엔진", "성능 병목 진단 (Force gather, Mass scaling, Contact, MPP 불균형)"],
            ["분석 엔진", "경고 패턴 분석 (tied contact 50135/50136)"],
            ["빌드", "Linux / Windows 빌드 스크립트 완성"],
            ["리포트", "Rich 터미널 / HTML / JSON 3개 형식 완성"],
        ],
        col_widths=[1.2, 5.3]
    )

    # ══════════════════════════════════════════════════════════════════════
    # 12. 현재 상태
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "12. 현재 상태 및 향후 계획", 1)

    add_heading(doc, "12.1 기능 완료 현황", 2)
    STATUS_COMPLETE = "완료"
    make_table(doc,
        ["#", "항목", "상태"],
        [
            ["1", "단일 폴더 분석 (터미널/HTML/JSON)", STATUS_COMPLETE],
            ["2", "HTML/JSON/터미널 리포트 (18개 섹션)", STATUS_COMPLETE],
            ["2a", "에너지 시계열 SVG 차트 (4종)", STATUS_COMPLETE],
            ["2b", "rcforc 접촉 반력 파싱 + SVG 차트", STATUS_COMPLETE],
            ["3", "error_db 68개 코드", STATUS_COMPLETE],
            ["4", "/data 인덱싱 (2,495개)", STATUS_COMPLETE],
            ["5", "/data 배치 분석 (2,440개)", STATUS_COMPLETE],
            ["6", "matsum 재료별 HG 분석", STATUS_COMPLETE],
            ["7", "Implicit solver 진단", STATUS_COMPLETE],
            ["8", "element_mapper 파트 보완", STATUS_COMPLETE],
            ["9", "INCOMPLETE 3분류 진단", STATUS_COMPLETE],
            ["10", "노이즈 제거 최적화", STATUS_COMPLETE],
            ["11", "Linux/Windows 빌드", STATUS_COMPLETE],
            ["12", "GUI 모드 (tkinter)", STATUS_COMPLETE],
        ],
        col_widths=[0.3, 4.0, 1.0]
    )

    add_heading(doc, "12.2 진단 커버리지 요약", 2)
    make_table(doc,
        ["메트릭", "값"],
        [
            ["에러 코드 DB", "68개"],
            ["파서 모듈", "11개"],
            ["분석 엔진", "10개"],
            ["SVG 차트", "8종"],
            ["진단 항목", "30+개"],
            ["검증 케이스", "2,440개"],
            ["ERROR → CRITICAL 커버리지", "100%"],
            ["INCOMPLETE → CRITICAL 커버리지", "100%"],
            ["WARNING+ 보유 비율", "95.8%"],
        ],
        col_widths=[3.0, 3.5]
    )

    add_heading(doc, "12.3 알려진 제한사항", 2)
    make_table(doc,
        ["#", "제한사항", "비고"],
        [
            ["1", "단위 테스트 없음", "데이터 기반 대규모 검증(2,440건)으로 대체"],
            ["2", "SMP 모드 미지원", "MPP만 지원 (대부분의 실무 환경)"],
            ["3", "nodout/bndout 스트리밍 최적화 미완", "대용량 파일 시 메모리 사용량 증가 가능"],
            ["4", "ALE/SPH 진단 제한적", "기본 코드 등록 완료, 심층 진단 미구현"],
            ["5", "Thermal-only 해석 심층 진단 미구현", "Implicit thermal 케이스는 INFO-only (정상)"],
        ],
        col_widths=[0.3, 2.5, 3.7]
    )

    add_heading(doc, "12.4 향후 개발 방향", 2)
    make_table(doc,
        ["우선순위", "항목", "설명"],
        [
            ["높음", "단위 테스트 (pytest)", "주요 파서·분석 모듈에 대한 자동화 테스트"],
            ["높음", "다중 케이스 비교 리포트", "Study 단위 집계, 트렌드 분석"],
            ["중간", "Thermal/Coupled 진단", "열 해석 특유의 수렴 문제 진단"],
            ["중간", "d3plot 요소 품질 분석", "KooD3plotReader 연동 시 구현"],
            ["낮음", "CI/CD 자동 빌드", "GitHub Actions 기반"],
            ["낮음", "SMP 모드 지원", "MPP 외 SMP 환경 파서 확장"],
        ],
        col_widths=[0.8, 2.0, 3.7]
    )

    doc.save("FEATURE_STATUS_REPORT.docx")
    print("FEATURE_STATUS_REPORT.docx 생성 완료")


if __name__ == "__main__":
    build_docx()
