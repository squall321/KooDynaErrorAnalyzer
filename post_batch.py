"""
배치 분석 완료 후 실행:
1. reports2/ 폴더에 다양한 에러 타입 대표 케이스 복사
2. FEATURE_STATUS_REPORT.docx 최종 업데이트
"""

import json
import shutil
from pathlib import Path
from collections import defaultdict


REPORTS_DIR = Path("/data/koodyna_reports")
REPORTS2_DIR = Path("reports2")
INDEX_PATH = Path.home() / ".koodyna" / "index.json"


# ── 1. 전체 JSON 리포트 수집 ──────────────────────────────────────────────────

def collect_all_reports(reports_dir: Path) -> list[dict]:
    """Gather all JSON reports under reports_dir."""
    results = []
    for json_file in sorted(reports_dir.rglob("*.json")):
        try:
            data = json.loads(json_file.read_text())
            data["_path"] = str(json_file)
            data["_study_type"] = json_file.parent.name
            data["_case"] = json_file.stem
            results.append(data)
        except Exception:
            pass
    return results


# ── 2. 에러 타입별 대표 케이스 선정 ──────────────────────────────────────────

def select_diverse_cases(reports: list[dict]) -> dict[str, list[dict]]:
    """
    각 카테고리별 대표 케이스를 1~2개씩 선정.
    선정 기준: 해당 카테고리 finding이 있으면서 CRITICAL 건수가 많은 것 우선.
    """
    categories = {
        "정상_종료":                lambda r: _termination(r) == "NORMAL",
        "Negative_Volume_40509":    lambda r: _has_error(r, 40509),
        "NaN_감지_40455":           lambda r: _has_error(r, 40455),
        "Segfault_Signal11":        lambda r: _has_title(r, "Segmentation fault"),
        "에너지_발산":               lambda r: _has_title(r, "Total energy is increasing"),
        "에너지_비율_폭주":          lambda r: _has_title(r, "에너지 비율 폭주"),
        "Hourglass_CRITICAL":       lambda r: _has_title(r, "Hourglass energy critically high"),
        "Hourglass_WARNING":        lambda r: _has_title(r, "Hourglass energy exceeds"),
        "NaN_에너지_glstat":        lambda r: _has_title(r, "NaN detected in energy"),
        "음수_Sliding":             lambda r: _has_title(r, "Negative sliding interface energy"),
        "KE_폭발":                  lambda r: _has_title(r, "Kinetic energy explosion"),
        "MPI_통신_에러":            lambda r: _has_title(r, "MPI communication error"),
        "MPP_불균형_CRITICAL":      lambda r: _has_title(r, "MPP load imbalance") and _has_crit_title(r, "MPP load imbalance"),
        "MPP_불균형_WARNING":       lambda r: _has_title(r, "MPP load imbalance") and not _has_crit_title(r, "MPP load imbalance"),
        "타임스텝_붕괴":            lambda r: _has_title(r, "Severe timestep drop"),
        "접촉_에너지_과다":         lambda r: _has_title(r, "Excessive contact sliding energy"),
        "Warning_40533":            lambda r: _has_warning_code(r, 40533),
        "Warning_40538":            lambda r: _has_warning_code(r, 40538),
        "Warning_21129":            lambda r: _has_warning_code(r, 21129),
        "Warning_40509_반복":       lambda r: _warning_count(r, 40509) >= 100,
        "에러_종료_exit3":          lambda r: _has_title(r, "LS-DYNA exit code 3"),
        "접촉_계산_과다":           lambda r: _has_title(r, "접촉 계산 시간 높음"),
        "Tied_Contact_50135":       lambda r: _has_warning_code(r, 50135),
        "Implicit_에러_60315":      lambda r: _has_error(r, 60315),
        "Curve_외삽_21129":         lambda r: _has_error(r, 21129),
    }

    selected: dict[str, list[dict]] = {}

    for category, matcher in categories.items():
        matches = [r for r in reports if matcher(r)]
        if not matches:
            continue
        # Sort by CRITICAL count descending, pick top 2
        matches.sort(key=lambda r: -_crit_count(r))
        selected[category] = matches[:2]

    return selected


def _termination(r):
    return r.get("termination", {}).get("status", "")

def _findings(r):
    return r.get("findings", [])

def _has_title(r, keyword):
    return any(keyword in f.get("title", "") for f in _findings(r))

def _has_crit_title(r, keyword):
    return any(
        keyword in f.get("title", "") and f.get("severity") == "CRITICAL"
        for f in _findings(r)
    )

def _has_error(r, code):
    warns = r.get("warnings", [])
    return any(w.get("code") == code for w in warns)

def _has_warning_code(r, code):
    warns = r.get("warnings", [])
    return any(w.get("code") == code and w.get("count", 0) > 0 for w in warns)

def _warning_count(r, code):
    warns = r.get("warnings", [])
    for w in warns:
        if w.get("code") == code:
            return w.get("count", 0)
    return 0

def _crit_count(r):
    return sum(1 for f in _findings(r) if f.get("severity") == "CRITICAL")


# ── 3. reports2 폴더 구성 ─────────────────────────────────────────────────────

def build_reports2(selected: dict[str, list[dict]], reports2_dir: Path):
    reports2_dir.mkdir(exist_ok=True)

    summary = []  # (category, case_name, study_type, crit, warn, findings_summary)

    for category, cases in sorted(selected.items()):
        for case in cases:
            src = Path(case["_path"])
            study = case["_study_type"]
            case_name = case["_case"]
            dest_name = f"{category}__{study}__{case_name}.json"
            dest = reports2_dir / dest_name
            shutil.copy2(src, dest)

            crits = [f for f in _findings(case) if f.get("severity") == "CRITICAL"]
            warns = [f for f in _findings(case) if f.get("severity") == "WARNING"]
            summary.append({
                "category": category,
                "case": f"{study}/{case_name}",
                "termination": _termination(case),
                "crit": len(crits),
                "warn": len(warns),
                "top_finding": crits[0]["title"] if crits else (warns[0]["title"] if warns else "-"),
            })

    # Write summary JSON
    summary_path = reports2_dir / "_summary.json"
    summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2))

    return summary


# ── 4. 전체 통계 ─────────────────────────────────────────────────────────────

def compute_global_stats(reports: list[dict]) -> dict:
    total = len(reports)
    by_status = defaultdict(int)
    by_failure = defaultdict(int)
    crit_total = warn_total = 0

    warning_code_counts: dict[int, int] = defaultdict(int)

    for r in reports:
        status = _termination(r)
        by_status[status] += 1
        crit_total += _crit_count(r)
        warn_total += sum(1 for f in _findings(r) if f.get("severity") == "WARNING")

        for f in _findings(r):
            title = f.get("title", "")
            if f.get("severity") == "CRITICAL":
                by_failure[title] += 1

        for w in r.get("warnings", []):
            code = w.get("code")
            cnt = w.get("count", 0)
            if code and cnt:
                warning_code_counts[code] += cnt

    top_failures = sorted(by_failure.items(), key=lambda x: -x[1])[:15]
    top_warnings = sorted(warning_code_counts.items(), key=lambda x: -x[1])[:15]

    return {
        "total_cases": total,
        "by_status": dict(by_status),
        "total_crits": crit_total,
        "total_warns": warn_total,
        "top_failure_findings": top_failures,
        "top_warning_codes": top_warnings,
    }


# ── 5. DOCX 업데이트 ──────────────────────────────────────────────────────────

def update_docx(stats: dict, summary: list[dict], n_reports2: int):
    """Append batch analysis results section to FEATURE_STATUS_REPORT.docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document("FEATURE_STATUS_REPORT.docx")

    def shade_row(row, hex_color="D9E1F2"):
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), hex_color)
            tcPr.append(shd)

    def make_table(headers, rows, col_widths=None):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_row = table.rows[0]
        shade_row(hdr_row, "2F5496")
        for i, h in enumerate(headers):
            run = hdr_row.cells[i].paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for ri, row_data in enumerate(rows):
            row = table.rows[ri + 1]
            if ri % 2 == 0:
                shade_row(row, "EEF2F9")
            for ci, val in enumerate(row_data):
                run = row.cells[ci].paragraphs[0].add_run(str(val))
                run.font.size = Pt(9)
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Inches(w)
        doc.add_paragraph()
        return table

    # Add new section
    doc.add_page_break()
    h = doc.add_heading("11. /data 배치 분석 결과 (2,407개 폴더)", 1)
    h.paragraph_format.space_before = Pt(10)

    # 11.1 전체 통계
    doc.add_heading("11.1 전체 분석 통계", 2)
    by_status = stats["by_status"]
    make_table(
        ["항목", "값"],
        [
            ["총 분석 케이스", f"{stats['total_cases']:,}개"],
            ["정상 종료 (NORMAL)", f"{by_status.get('NORMAL', 0):,}개"],
            ["에러 종료 (ERROR)", f"{by_status.get('ERROR', 0):,}개"],
            ["미완료 (INCOMPLETE)", f"{by_status.get('INCOMPLETE', 0):,}개"],
            ["기타 / 미분류", f"{by_status.get('', 0) + by_status.get('UNKNOWN', 0):,}개"],
            ["총 CRITICAL 진단 건수", f"{stats['total_crits']:,}건"],
            ["총 WARNING 진단 건수", f"{stats['total_warns']:,}건"],
        ],
        col_widths=[3.0, 3.0]
    )

    # 11.2 상위 실패 원인
    doc.add_heading("11.2 상위 CRITICAL 진단 패턴 (전체 케이스)", 2)
    make_table(
        ["CRITICAL Finding 제목", "발생 케이스 수"],
        [(title, f"{cnt:,}") for title, cnt in stats["top_failure_findings"]],
        col_widths=[4.5, 1.5]
    )

    # 11.3 상위 경고 코드
    doc.add_heading("11.3 상위 Warning/Error 코드 (누적 발생 횟수)", 2)
    make_table(
        ["코드", "누적 발생 횟수"],
        [(f"Warning/Error {code}", f"{cnt:,}") for code, cnt in stats["top_warning_codes"]],
        col_widths=[2.5, 3.5]
    )

    # 11.4 reports2 선정 케이스
    doc.add_heading(f"11.4 다양성 검증 케이스 (reports2/ — {n_reports2}개)", 2)
    p = doc.add_paragraph(
        "각 에러 유형별 대표 케이스를 reports2/ 폴더에 수집했습니다. "
        "파일명 형식: <카테고리>__<study_type>__<case_name>.json"
    )
    p.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    make_table(
        ["카테고리", "케이스", "종료", "CRIT", "WARN", "대표 Finding"],
        [
            [
                s["category"], s["case"], s["termination"],
                s["crit"], s["warn"], s["top_finding"][:55]
            ]
            for s in summary
        ],
        col_widths=[1.5, 1.5, 0.8, 0.4, 0.4, 2.4]
    )

    doc.save("FEATURE_STATUS_REPORT.docx")
    print("FEATURE_STATUS_REPORT.docx 업데이트 완료")


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("리포트 수집 중...")
    reports = collect_all_reports(REPORTS_DIR)
    # also include the 10 test cases in local reports/
    local_reports = []
    for jf in sorted(Path("reports").glob("*.json")):
        try:
            data = json.loads(jf.read_text())
            data["_path"] = str(jf)
            data["_study_type"] = "results"
            data["_case"] = jf.stem
            local_reports.append(data)
        except Exception:
            pass
    all_reports = reports + local_reports
    print(f"  총 {len(all_reports):,}개 리포트 수집")

    print("다양성 케이스 선정 중...")
    selected = select_diverse_cases(all_reports)
    print(f"  {len(selected)}개 카테고리, {sum(len(v) for v in selected.values())}개 케이스 선정")

    print(f"reports2/ 구성 중...")
    summary = build_reports2(selected, REPORTS2_DIR)
    print(f"  {len(summary)}개 파일 복사 완료 → {REPORTS2_DIR}/")

    print("전체 통계 계산 중...")
    stats = compute_global_stats(all_reports)
    print(f"  정상: {stats['by_status'].get('NORMAL',0)}, "
          f"에러: {stats['by_status'].get('ERROR',0)}, "
          f"미완료: {stats['by_status'].get('INCOMPLETE',0)}")
    print(f"  총 CRITICAL: {stats['total_crits']:,}, WARNING: {stats['total_warns']:,}")

    print("DOCX 업데이트 중...")
    update_docx(stats, summary, len(summary))

    print("\n완료!")
    print(f"  reports2/ : {len(summary)}개 케이스")
    print(f"  FEATURE_STATUS_REPORT.docx : 섹션 11 추가됨")
